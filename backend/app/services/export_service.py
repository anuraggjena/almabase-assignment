import io
from docx import Document
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from reportlab.lib.units import inch
from reportlab.lib.styles import getSampleStyleSheet
from reportlab.lib.pagesizes import letter
from reportlab.lib import utils
from reportlab.lib.styles import ParagraphStyle
from reportlab.lib import colors
from xml.sax.saxutils import escape


# docx export

def export_docx(questions, answers):

    doc = Document()

    for q in questions:
        doc.add_heading(f"{q['order'] + 1}. {q['text']}", level=2)

        answer_obj = next(
            (a for a in answers if a["question_id"] == q["id"]),
            None
        )

        if answer_obj:
            doc.add_paragraph("Answer:")
            doc.add_paragraph(answer_obj["answer"])

            doc.add_paragraph(
                f"Confidence Score: {answer_obj['confidence_score']}%"
            )

            if answer_obj.get("citations"):
                doc.add_paragraph("Evidence:")
                for c in answer_obj["citations"]:
                    doc.add_paragraph(
                        f"{c.get('document_name','Reference')}, "
                        f"Chunk {c['chunk_index']}: "
                        f"{c.get('snippet','')}"
                    )

        doc.add_paragraph("")
        doc.add_paragraph("-" * 60)
        doc.add_paragraph("")

    file_stream = io.BytesIO()
    doc.save(file_stream)
    file_stream.seek(0)

    return file_stream

# txt export

def export_txt(questions, answers):

    content = []

    for q in questions:
        content.append(f"{q['order'] + 1}. {q['text']}\n")

        answer_obj = next(
            (a for a in answers if a["question_id"] == q["id"]),
            None
        )

        if answer_obj:
            content.append("Answer:\n")
            content.append(f"{answer_obj['answer']}\n\n")
            content.append(
                f"Confidence Score: {answer_obj['confidence_score']}%\n"
            )

            if answer_obj.get("citations"):
                content.append("Evidence:\n")
                for c in answer_obj["citations"]:
                    content.append(
                        f"{c.get('document_name','Reference')}, "
                        f"Chunk {c['chunk_index']}: "
                        f"{c.get('snippet','')}\n"
                    )

        content.append("\n" + "-" * 60 + "\n\n")

    final_text = "".join(content)

    return io.BytesIO(final_text.encode("utf-8"))

# pdf export

def export_pdf(questions, answers):

    buffer = io.BytesIO()
    doc = SimpleDocTemplate(buffer, pagesize=letter)
    elements = []

    styles = getSampleStyleSheet()

    for q in questions:

        elements.append(
            Paragraph(
                f"<b>{q['order'] + 1}. {escape(q['text'])}</b>",
                styles["Normal"]
            )
        )
        elements.append(Spacer(1, 0.2 * inch))

        answer_obj = next(
            (a for a in answers if a["question_id"] == q["id"]),
            None
        )

        if answer_obj:

            elements.append(
                Paragraph("<b>Answer:</b>", styles["Normal"])
            )
            elements.append(Spacer(1, 0.1 * inch))

            elements.append(
                Paragraph(escape(answer_obj["answer"]), styles["Normal"])
            )
            elements.append(Spacer(1, 0.2 * inch))

            elements.append(
                Paragraph(
                    f"<b>Confidence Score:</b> "
                    f"{answer_obj['confidence_score']}%",
                    styles["Normal"]
                )
            )
            elements.append(Spacer(1, 0.2 * inch))

            if answer_obj.get("citations"):
                elements.append(
                    Paragraph("<b>Evidence:</b>", styles["Normal"])
                )
                elements.append(Spacer(1, 0.1 * inch))

                for c in answer_obj["citations"]:
                    elements.append(
                        Paragraph(
                            f"{escape(c.get('document_name','Reference'))}, "
                            f"Chunk {c['chunk_index']}: "
                            f"{escape(c.get('snippet',''))}",
                            styles["Normal"]
                        )
                    )
                    elements.append(Spacer(1, 0.2 * inch))

        elements.append(Spacer(1, 0.5 * inch))

    doc.build(elements)
    buffer.seek(0)

    return buffer